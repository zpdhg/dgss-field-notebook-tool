#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""分册合并脚本"""
import os, glob, re, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
from PIL import Image
import io

def extract_route_number(filename):
    match = re.search(r'L(\d+)', filename)
    return (int(match.group(1)), match.group(0)) if match else (None, None)

def get_sorted_files(directory, pattern):
    files = glob.glob(os.path.join(directory, pattern))
    file_info = []
    for filename in files:
        basename = os.path.basename(filename)
        route_num, route_name = extract_route_number(basename)
        if route_num is not None:
            file_info.append((route_num, route_name, filename))
    file_info.sort(key=lambda x: x[0])
    return file_info

def set_images_dpi(doc_path, dpi=330):
    try:
        doc = Document(doc_path)
        for shape in doc.inline_shapes:
            rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = doc.part.related_parts[rId]
            img = Image.open(io.BytesIO(image_part.blob))
            output = io.BytesIO()
            img.save(output, format=img.format, dpi=(dpi, dpi))
            image_part._blob = output.getvalue()
        doc.save(doc_path)
    except: pass

def ensure_document_layout(doc):
    """[FIXED] 确保文档布局正确：主要内容双栏，素描图部分单栏。
    使用精确的节定位，并且只将包含素描图标题的节设为单栏。
    """
    try:
        import re
        
        # 1. 默认设置所有节为双栏
        print("开始调整文档布局...")
        for i, section in enumerate(doc.sections):
            try:
                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')
                if cols:
                    cols[0].set(qn('w:num'), '2')
                else:
                    cols = OxmlElement('w:cols')
                    cols.set(qn('w:num'), '2')
                    sectPr.append(cols)
            except Exception as e:
                print(f"设置节 {i+1} 布局为双栏时出错: {e}")
                continue
        print("所有节已默认设置为双栏。")
        
        # 2. 查找所有包含素描图标题的段落，并记录它们所属的节索引
        sketch_sections_to_update = set()
        
        for i, para in enumerate(doc.paragraphs):
            # 匹配 "LXXXX素描图" 格式
            if re.match(r'^L\d+素描图$', para.text.strip()):
                print(f"发现素描图标题: '{para.text}' 在段落 {i}")
                
                # 使用修复后的函数查找该段落所属的节
                section_index = find_paragraph_section(doc, i)
                if section_index is not None:
                    print(f"-> 素描图标题属于第 {section_index + 1} 节")
                    sketch_sections_to_update.add(section_index)
                else:
                    print(f"-> 未能确定素描图标题所属的节")

        # 3. 将所有记录的需要更新的节设置为单栏
        print(f"\n总共有 {len(sketch_sections_to_update)} 个节需要设置为单栏。")
        for section_index in sorted(list(sketch_sections_to_update)):
            try:
                if section_index < len(doc.sections):
                    section = doc.sections[section_index]
                    sectPr = section._sectPr
                    cols = sectPr.xpath('./w:cols')
                    if cols:
                        cols[0].set(qn('w:num'), '1')
                        print(f"  ✓ 已将第 {section_index + 1} 节设置为单栏")
                    else:
                        cols = OxmlElement('w:cols')
                        cols.set(qn('w:num'), '1')
                        sectPr.append(cols)
                        print(f"  ✓ 已为第 {section_index + 1} 节创建并设置为单栏")
                else:
                    print(f"  ✗ 警告：计算出的节索引 {section_index} 超出范围 (总节数: {len(doc.sections)})")
            except Exception as e:
                print(f"设置第 {section_index + 1} 节为单栏时出错: {e}")
                    
        print("\n文档布局自动调整完成。")
                
    except Exception as e:
        print(f"确保文档布局时发生严重错误: {e}")

def find_paragraph_section(doc, paragraph_index):
    """[FIXED] 查找指定段落索引所属的节索引。
    通过分析文档结构中实际的节定义（sectPr）来准确定位，修复了原先错误的简化实现。
    """
    try:
        # 1. 构建一个映射，记录每个节在哪个段落之后结束。
        # sectPr 元素定义了其“之前”文本的节的属性。
        # doc.paragraphs 是一个扁平列表，我们需要找到其中的分节符。
        section_end_para_map = {}
        current_section_idx = 0
        for i, p in enumerate(doc.paragraphs):
            # 如果段落属性包含节属性，说明这里是一个分节符
            if p._p.pPr and p._p.pPr.sectPr:
                section_end_para_map[current_section_idx] = i
                current_section_idx += 1
        
        # 2. 根据段落索引在哪个节的范围内来确定其所属的节。
        # 我们需要对节的结束位置进行排序，以便正确查找。
        # 假定第一个节从段落0开始。
        last_para_idx = -1
        for sec_idx in sorted(section_end_para_map.keys()):
            end_para_idx = section_end_para_map[sec_idx]
            # 检查目标段落是否在本节范围内
            if last_para_idx < paragraph_index <= end_para_idx:
                return sec_idx
            last_para_idx = end_para_idx
            
        # 如果段落在所有显式分节符之后，那么它属于最后一个节
        return current_section_idx
        
    except Exception as e:
        print(f"查找段落节索引时出错 (新方法): {e}")
        return None

def add_page_number(run):
    """为run添加页码域 (PAGE field) 的底层XML。"""
    # 设置字体为宋体
    run.font.name = '宋体'
    run.font.size = Pt(9)
    if run._element.rPr is None:
        run._element._add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def enforce_header_styles(doc):
    """强制设置 RouteHeader 和 SectionHeader 的字体颜色为黑色"""
    print("正在强制设置标题样式颜色为黑色...")
    
    # 1. 尝试直接修改样式定义
    target_styles = ['RouteHeader', 'SectionHeader', '标题 1', '标题 2', 'Heading 1', 'Heading 2']
    for style_name in target_styles:
        if style_name in doc.styles:
            try:
                style = doc.styles[style_name]
                style.font.color.rgb = RGBColor(0, 0, 0)
                print(f"  ✓ 已设置样式 '{style_name}' 颜色为黑色")
            except Exception as e:
                print(f"  ✗ 设置样式 '{style_name}' 失败: {e}")

    # 2. 遍历段落，针对特定样式强制设置
    # 这一步是为了处理那些虽然应用了样式但可能被手动覆盖格式的段落
    count = 0
    for para in doc.paragraphs:
        if para.style.name in target_styles:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            count += 1
    print(f"  ✓ 已检查并修正 {count} 个标题段落的颜色")

def format_sketch_map_titles(doc):
    """强制设置素描图标题的格式：宋体、五号(10.5pt)、左对齐"""
    print("正在强制设置素描图标题格式...")
    count = 0
    for para in doc.paragraphs:
        # 匹配 "LXXXX素描图" 格式
        if re.match(r'^L\d+素描图$', para.text.strip()):
            # 设置对齐方式
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 设置字体
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                if run._element.rPr is None:
                    run._element._add_rPr()
                rFonts = run._element.rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:ascii'), '宋体')
                rFonts.set(qn('w:hAnsi'), '宋体')
            count += 1
    print(f"  ✓ 已修正 {count} 个素描图标题的格式")

def merge_volume(file_info_list, volume_num, output_dir):
    if not file_info_list:
        return None
    
    # 1. 先将所有路线报告合并到一个临时文档中
    master_path = file_info_list[0][2]
    content_composer = Composer(Document(master_path))
    for i, (_, _, filename) in enumerate(file_info_list[1:], 1):
        try:
            content_composer.doc.add_section(WD_SECTION_START.NEW_PAGE)
            content_composer.append(Document(filename))
        except Exception as e:
            print(f"合并文档 {filename} 时出错: {e}")
            continue
    
    # 2. 创建一个包含封面和正确页码设置的新文档
    final_doc = Document()
    # 封面和空白页所在的节（第一节）
    # 确保第一节没有页码
    section1 = final_doc.sections[0]
    section1.footer.is_linked_to_previous = False
    # 清空第一节页脚（如果有的话）
    if section1.footer.paragraphs:
        for p in section1.footer.paragraphs:
            p.clear()
            
    page1 = final_doc.add_paragraph()
    run1 = page1.add_run('封面')
    run1.font.size = Pt(28)
    run1.font.bold = True
    page1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page1.paragraph_format.space_before = Pt(200)
    final_doc.add_paragraph().paragraph_format.page_break_before = True
    final_doc.add_paragraph().paragraph_format.page_break_before = True
    final_doc.add_paragraph().paragraph_format.page_break_before = True
    
    # 3. 添加一个新节，用于存放主内容，并设置页码从1开始
    main_section = final_doc.add_section(WD_SECTION_START.NEW_PAGE)
    sectPr = main_section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')  # 页码从1开始
    sectPr.append(pgNumType)

    # 4. 在新节的页脚添加页码
    # 关键：取消链接到前一节，确保页码只出现在本节及后续
    main_section.footer.is_linked_to_previous = False
    
    footer = main_section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p.add_run())

    # 5. 将封面文档与内容文档合并
    final_composer = Composer(final_doc)
    final_composer.append(content_composer.doc)

    first_route = file_info_list[0][1]
    last_route = file_info_list[-1][1]
    output_filename = os.path.join(output_dir, f"野外手图_第{volume_num}册_{first_route}-{last_route}.docx")
    
    try:
        final_composer.save(output_filename)
        
        # 确保文档布局正确
        doc_to_process = Document(output_filename)
        ensure_document_layout(doc_to_process)
        
        # 强制设置标题颜色
        enforce_header_styles(doc_to_process)
        
        # 强制设置素描图标题格式
        format_sketch_map_titles(doc_to_process)
        
        doc_to_process.save(output_filename)
        
        # 设置图片DPI
        set_images_dpi(output_filename, dpi=330)
        return output_filename
    except Exception as e:
        print(f"保存最终文档 {output_filename} 时出错: {e}")
        return None

def run_batch():
    script_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
    input_dir = os.path.join(script_dir, "报告-已插入素描图")
    output_dir = os.path.join(script_dir, "分册")
    
    if not os.path.exists(input_dir): 
        print(f"输入目录不存在: {input_dir}")
        return
    if not os.path.exists(output_dir): 
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 获取所有文件，按路线序号顺序分册，不分版本类型
    all_files = []
    all_files.extend(get_sorted_files(input_dir, "L*完整版.docx"))
    all_files.extend(get_sorted_files(input_dir, "L*_formatted.docx"))
    
    # 按路线编号去重，优先使用完整版
    unique_files = {}
    for route_num, route_name, filepath in all_files:
        if route_num not in unique_files:
            unique_files[route_num] = (route_num, route_name, filepath)
        else:
            # 如果已有相同路线编号，优先选择完整版
            existing_filepath = unique_files[route_num][2]
            if "完整版" in filepath and "formatted" in existing_filepath:
                unique_files[route_num] = (route_num, route_name, filepath)
    
    # 按路线编号排序
    file_info = sorted(unique_files.values(), key=lambda x: x[0])
    
    if not file_info: 
        print("未找到任何可处理的文件")
        return
    
    print(f"找到 {len(file_info)} 个路线文件，按编号排序:")
    for route_num, route_name, _ in file_info:
        print(f"  {route_name} (编号: {route_num})")
    
    routes_per_volume = 12
    total_routes = len(file_info)
    total_volumes = (total_routes + routes_per_volume - 1) // routes_per_volume
    
    print(f"\n将分成 {total_volumes} 册，每册最多 {routes_per_volume} 个路线")
    
    for volume_num in range(1, total_volumes + 1):
        start_idx = (volume_num - 1) * routes_per_volume
        end_idx = min(start_idx + routes_per_volume, total_routes)
        volume_files = file_info[start_idx:end_idx]
        
        print(f"\n处理第 {volume_num} 册: {volume_files[0][1]} 到 {volume_files[-1][1]}")
        result = merge_volume(volume_files, volume_num, output_dir)
        if result:
            print(f"✓ 成功创建: {os.path.basename(result)}")
        else:
            print(f"✗ 第 {volume_num} 册创建失败")

def run_batch_with_routes_per_volume(routes_per_volume):
    """使用自定义每册路线数运行分册合并"""
    script_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
    input_dir = os.path.join(script_dir, "报告-已插入素描图")
    output_dir = os.path.join(script_dir, "分册")
    
    if not os.path.exists(input_dir): 
        print(f"输入目录不存在: {input_dir}")
        return
    if not os.path.exists(output_dir): 
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 获取所有文件，按路线序号顺序分册，不分版本类型
    all_files = []
    all_files.extend(get_sorted_files(input_dir, "L*完整版.docx"))
    all_files.extend(get_sorted_files(input_dir, "L*_formatted.docx"))
    
    # 按路线编号去重，优先使用完整版
    unique_files = {}
    for route_num, route_name, filepath in all_files:
        if route_num not in unique_files:
            unique_files[route_num] = (route_num, route_name, filepath)
        else:
            # 如果已有相同路线编号，优先选择完整版
            existing_filepath = unique_files[route_num][2]
            if "完整版" in filepath and "formatted" in existing_filepath:
                unique_files[route_num] = (route_num, route_name, filepath)
    
    # 按路线编号排序
    file_info = sorted(unique_files.values(), key=lambda x: x[0])
    
    if not file_info: 
        print("未找到任何可处理的文件")
        return
    
    print(f"找到 {len(file_info)} 个路线文件，按编号排序:")
    for route_num, route_name, _ in file_info:
        print(f"  {route_name} (编号: {route_num})")
    
    total_routes = len(file_info)
    total_volumes = (total_routes + routes_per_volume - 1) // routes_per_volume
    
    print(f"\n每册 {routes_per_volume} 条路线，总共分为 {total_volumes} 册")
    
    for volume_num in range(1, total_volumes + 1):
        start_idx = (volume_num - 1) * routes_per_volume
        end_idx = min(start_idx + routes_per_volume, total_routes)
        volume_files = file_info[start_idx:end_idx]
        
        print(f"\n处理第 {volume_num} 册: {volume_files[0][1]} 到 {volume_files[-1][1]}")
        result = merge_volume(volume_files, volume_num, output_dir)
        if result:
            print(f"✓ 成功创建: {os.path.basename(result)}")
        else:
            print(f"✗ 第 {volume_num} 册创建失败")

def run_batch_with_total_volumes(total_volumes):
    """使用自定义总册数运行分册合并"""
    script_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
    input_dir = os.path.join(script_dir, "报告-已插入素描图")
    output_dir = os.path.join(script_dir, "分册")
    
    if not os.path.exists(input_dir): 
        print(f"输入目录不存在: {input_dir}")
        return
    if not os.path.exists(output_dir): 
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 获取所有文件，按路线序号顺序分册，不分版本类型
    all_files = []
    all_files.extend(get_sorted_files(input_dir, "L*完整版.docx"))
    all_files.extend(get_sorted_files(input_dir, "L*_formatted.docx"))
    
    # 按路线编号去重，优先使用完整版
    unique_files = {}
    for route_num, route_name, filepath in all_files:
        if route_num not in unique_files:
            unique_files[route_num] = (route_num, route_name, filepath)
        else:
            # 如果已有相同路线编号，优先选择完整版
            existing_filepath = unique_files[route_num][2]
            if "完整版" in filepath and "formatted" in existing_filepath:
                unique_files[route_num] = (route_num, route_name, filepath)
    
    # 按路线编号排序
    file_info = sorted(unique_files.values(), key=lambda x: x[0])
    
    if not file_info: 
        print("未找到任何可处理的文件")
        return
    
    print(f"找到 {len(file_info)} 个路线文件，按编号排序:")
    for route_num, route_name, _ in file_info:
        print(f"  {route_name} (编号: {route_num})")
    
    total_routes = len(file_info)
    routes_per_volume = total_routes // total_volumes
    if total_routes % total_volumes != 0:
        routes_per_volume += 1  # 向上取整
    
    print(f"\n总共分为 {total_volumes} 册，每册约 {routes_per_volume} 条路线")
    
    for volume_num in range(1, total_volumes + 1):
        start_idx = (volume_num - 1) * routes_per_volume
        end_idx = min(start_idx + routes_per_volume, total_routes)
        volume_files = file_info[start_idx:end_idx]
        
        if volume_files:
            print(f"\n处理第 {volume_num} 册: {volume_files[0][1]} 到 {volume_files[-1][1]}")
            result = merge_volume(volume_files, volume_num, output_dir)
            if result:
                print(f"✓ 成功创建: {os.path.basename(result)}")
            else:
                print(f"✗ 第 {volume_num} 册创建失败")

if __name__ == "__main__":
    run_batch()