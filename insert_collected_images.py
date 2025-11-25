#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
步骤 3: 插入已汇总的素描图
读取格式化后的报告和“素描图汇总”文件夹中的图片，
将图片插入报告末尾，并输出到“报告-已插入素描图”文件夹。
"""

import os
import sys
import glob
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START

def get_script_dir():
    """获取脚本或exe所在的目录"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

def extract_route_name(filename):
    """从文件名中提取路线名称 (e.g., L0459)"""
    match = re.search(r'(L\d+)', os.path.basename(filename))
    if match:
        return match.group(1)
    return None

def set_section_columns(section, num_columns):
    """设置节的分栏数（兼容性更好的方法）"""
    try:
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), str(num_columns))
        else:
            cols = OxmlElement('w:cols')
            cols.set(qn('w:num'), str(num_columns))
            sectPr.append(cols)
    except Exception as e:
        print(f"    警告：设置分栏失败: {e}")

def run_batch():
    """批量将汇总的图片插入到格式化报告中"""
    script_dir = get_script_dir()
    
    input_reports_dir = os.path.join(script_dir, "重新排版的报告")
    input_images_dir = os.path.join(script_dir, "素描图汇总")
    output_dir = os.path.join(script_dir, "报告-已插入素描图")

    print("=" * 60)
    print("步骤 3: 插入素描图")
    print("=" * 60)
    print(f"读取格式化报告位置: {input_reports_dir}")
    print(f"读取素描图汇总位置: {input_images_dir}")
    print(f"插图后报告输出位置: {output_dir}")

    # 检查输入文件夹
    if not os.path.exists(input_reports_dir):
        print(f"\n错误：未找到格式化报告文件夹 '{os.path.basename(input_reports_dir)}'。请先运行步骤1。")
        return
    if not os.path.exists(input_images_dir):
        print(f"\n错误：未找到素描图汇总文件夹 '{os.path.basename(input_images_dir)}'。请先运行步骤2。")
        return
    
    # 创建输出文件夹
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出文件夹: {output_dir}")

    # 1. 收集所有格式化报告
    report_files = glob.glob(os.path.join(input_reports_dir, "*_formatted.docx"))
    reports_map = {extract_route_name(f): f for f in report_files if extract_route_name(f)}

    # 2. 收集所有汇总图片
    image_files = glob.glob(os.path.join(input_images_dir, "*.png"))
    images_map = {}
    for img in image_files:
        route_name = extract_route_name(img)
        if route_name:
            if route_name not in images_map:
                images_map[route_name] = []
            images_map[route_name].append(img)
    
    # 按数字后缀排序图片
    for route_name in images_map:
        images_map[route_name].sort()

    print(f"\n找到 {len(reports_map)} 个格式化报告和 {len(images_map)} 个路线的素描图。")
    
    processed_count = 0
    error_count = 0

    # 3. 遍历报告，插入图片
    for route_name, report_path in reports_map.items():
        print(f"\n处理路线: {route_name}")
        
        # 复制不需要插图的文件，并重命名
        if route_name not in images_map:
            try:
                # 文件名改为_完整版以统一格式，方便后续合并
                new_filename = os.path.join(output_dir, f"{route_name}_完整版.docx")
                import shutil
                shutil.copy2(report_path, new_filename)
                print(f"  - 未找到素描图，直接复制文件到: {os.path.basename(new_filename)}")
                processed_count += 1
            except Exception as e:
                print(f"  ✗ 复制文件时出错: {e}")
                error_count += 1
            continue

        try:
            doc = Document(report_path)
            sketch_images = images_map[route_name]
            print(f"  - 找到 {len(sketch_images)} 张素描图，准备插入...")

            # 添加新节，用于存放素描图，并设置为单栏
            new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            set_section_columns(new_section, 1)

            # 添加大标题 "Lxxxx素描图"
            title_text = f"{route_name}素描图"
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title_text)
            # 设置左对齐
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 设置字体: 宋体, 五号(10.5pt)
            font = title_run.font
            font.name = '宋体'
            font.size = Pt(10.5)
            font.bold = True
            
            # 强制设置中文字体
            if title_run._element.rPr is None:
                title_run._element._add_rPr()
            rFonts = title_run._element.rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), '宋体')
            rFonts.set(qn('w:hAnsi'), '宋体')
            
            # 插入图片
            for img_path in sketch_images:
                try:
                    # 图片前添加空行
                    doc.add_paragraph()
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.add_run().add_picture(img_path, width=Inches(6.0))
                    print(f"    ✓ 成功插入图片: {os.path.basename(img_path)}")
                except Exception as e:
                    print(f"    ✗ 插入图片失败: {os.path.basename(img_path)} - {e}")
            
            # 保存到输出目录
            output_filename = os.path.join(output_dir, f"{route_name}_完整版.docx")
            doc.save(output_filename)
            print(f"  ✓ 已保存: {os.path.basename(output_filename)}")
            processed_count += 1

        except Exception as e:
            print(f"  ✗ 处理报告 {os.path.basename(report_path)} 时发生严重错误: {e}")
            error_count += 1
            continue
    
    print("\n" + "=" * 60)
    print("插图处理完成!")
    print(f"成功: {processed_count} 个文件")
    print(f"失败: {error_count} 个文件")
    print(f"输出位置: {output_dir}")
    print("=" * 60)

if __name__ == "__main__":
    run_batch()
