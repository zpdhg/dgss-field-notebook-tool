import docx
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def set_columns(doc, num_columns=2):
    """Set the document to have a specific number of columns.
    
    Includes error handling for different Word versions.
    """
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), str(num_columns))
        else:
            cols = OxmlElement('w:cols')
            cols.set(qn('w:num'), str(num_columns))
            sectPr.append(cols)
    except Exception as e:
        print(f"Warning: Could not set columns: {e}")

def format_paragraph_fonts(paragraph):
    """Set fonts to Times New Roman (Western) and SimSum (Chinese), size 10.5pt.
    
    Includes error handling for font operations.
    """
    for run in paragraph.runs:
        try:
            run.font.name = 'Times New Roman'
            # 确保rPr元素存在
            if run._element.rPr is None:
                run._element._add_rPr()
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)
        except Exception as e:
            print(f"Warning: Error formatting font in run: {e}")
            continue

def format_chunk_spacing(chunk):
    """Set line spacing to Single (1.0) for all paragraphs in a chunk."""
    for elem in chunk["elements"]:
        if elem.tag.endswith('p'):
            pPr = elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                elem.insert(0, pPr)
            
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            
            # Set spacing to Single (240)
            spacing.set(qn('w:line'), "240")
            spacing.set(qn('w:lineRule'), "auto")
            spacing.set(qn('w:after'), "0")

def resize_images(doc):
    """Resize images to fit within a column (approx 6.5cm).
    
    Includes error handling for image operations.
    """
    try:
        max_width = Cm(6.5)
        for shape in doc.inline_shapes:
            try:
                if shape.width > max_width:
                    aspect_ratio = shape.height / shape.width
                    shape.width = max_width
                    shape.height = int(max_width * aspect_ratio)
            except Exception as e:
                print(f"Warning: Could not resize image: {e}")
                continue
    except Exception as e:
        print(f"Warning: Error accessing inline shapes: {e}")

def get_text_from_element(element):
    """Extract text from a paragraph XML element."""
    if element.tag.endswith('p'):
        text = ""
        for node in element.iter(qn('w:t')):
            if node.text:
                text += node.text
        return text
    return ""

def replace_text_in_element(element, old_text, new_text):
    """Replace text in a paragraph element."""
    if element.tag.endswith('p'):
        for node in element.iter(qn('w:t')):
            if node.text and old_text in node.text:
                node.text = node.text.replace(old_text, new_text)

def normalize_colons_in_element(element):
    """Normalize all colons to Chinese colon (：) and remove extra spaces after colons."""
    if element.tag.endswith('p'):
        for node in element.iter(qn('w:t')):
            if node.text:
                # Replace English colon (with or without space) with Chinese colon (no space)
                node.text = node.text.replace(': ', '：')  # English colon + space
                node.text = node.text.replace(':', '：')   # English colon without space
                # Remove extra space after Chinese colon
                node.text = node.text.replace('： ', '：')  # Chinese colon + space

def get_base_style(doc, style_names):
    """Try to find a valid base style from a list of names. If not found, create the first one.
    
    This function is designed to be robust across different Word versions and languages.
    It attempts multiple fallback strategies to ensure a valid style is always returned.
    """
    styles = doc.styles
    
    # First attempt: Try to find an existing style from the provided list
    for name in style_names:
        try:
            if name in styles:
                return styles[name]
        except Exception as e:
            print(f"Warning: Error checking style '{name}': {e}")
            continue
            
    # Second attempt: Create the first style name from the list
    target_name = style_names[0]
    try:
        print(f"Style '{target_name}' not found. Creating it...")
        new_style = styles.add_style(target_name, 1)  # 1 = PARAGRAPH
        
        # Try to set a base style
        if 'Normal' in styles:
            try:
                new_style.base_style = styles['Normal']
            except:
                pass
        elif '正文' in styles:
            try:
                new_style.base_style = styles['正文']
            except:
                pass
        
        print(f"Successfully created style '{target_name}'")
        return new_style
    except Exception as e:
        print(f"Warning: Could not create base style '{target_name}': {e}")
    
    # Third attempt: Return Normal or 正文 style as fallback
    fallback_styles = ['Normal', '正文', 'Body Text', 'Body', 'Default Paragraph Font']
    for fallback in fallback_styles:
        try:
            if fallback in styles:
                print(f"Using fallback style '{fallback}'")
                return styles[fallback]
        except Exception as e:
            print(f"Warning: Error accessing fallback style '{fallback}': {e}")
            continue
    
    # Final attempt: Return the first available paragraph style
    try:
        for style in styles:
            if style.type.name == 'PARAGRAPH':
                print(f"Using first available paragraph style: '{style.name}'")
                return style
    except Exception as e:
        print(f"Warning: Error iterating styles: {e}")
    
    # Last resort: This should never happen, but return None with a clear error message
    print("CRITICAL: Could not find or create any valid style. Document may have issues.")
    return None

def detect_heading_style_names(doc):
    """检测中英文标题样式，优先使用中文版"""
    styles = doc.styles
    
    # 优先检测中文样式，然后英文样式
    heading1_variants = [
        '标题 1',        # Chinese Word (优先)
        'Heading 1',     # English Word
        '標題 1',        # Traditional Chinese
        'Titre 1',       # French
        'Überschrift 1', # German
        'Título 1',      # Spanish/Portuguese
    ]
    
    heading2_variants = [
        '标题 2',        # Chinese Word (优先)
        'Heading 2',     # English Word
        '標題 2',
        'Titre 2',
        'Überschrift 2',
        'Título 2',
    ]
    
    # 检测Heading 1
    detected_h1 = None
    detected_h2 = None
    
    # 先检查所有可用的样式名称（用于调试）
    available_heading_styles = []
    for style in styles:
        if style.type.name == 'PARAGRAPH':
            style_name = style.name
            if any(h in style_name for h in ['标题 1', 'Heading 1', '标题 2', 'Heading 2']):
                available_heading_styles.append(style_name)
    
    if available_heading_styles:
        print(f"Available heading styles: {available_heading_styles}")
    
    # 查找Heading 1
    for h1_name in heading1_variants:
        try:
            if h1_name in styles:
                detected_h1 = h1_name
                print(f"Detected Heading 1 style as: '{h1_name}'")
                break
        except Exception as e:
            print(f"Warning: Error checking style '{h1_name}': {e}")
            continue
    
    # 查找Heading 2
    for h2_name in heading2_variants:
        try:
            if h2_name in styles:
                detected_h2 = h2_name
                print(f"Detected Heading 2 style as: '{h2_name}'")
                break
        except Exception as e:
            print(f"Warning: Error checking style '{h2_name}': {e}")
            continue
    
    # 如果没有找到，使用默认值
    if not detected_h1:
        print("No heading 1 style detected, defaulting to '标题 1'")
        detected_h1 = '标题 1'
    
    if not detected_h2:
        print("No heading 2 style detected, defaulting to '标题 2'")
        detected_h2 = '标题 2'
    
    return detected_h1, detected_h2

def ensure_heading_styles_exist(doc):
    """确保Word文档中存在标题1和标题2样式，避免重复创建和自动重命名"""
    styles = doc.styles
    
    # 检查所有现有样式，包括可能的重复样式
    existing_styles = []
    for style in styles:
        if style.type.name == 'PARAGRAPH':
            existing_styles.append(style.name)
    
    print(f"Existing paragraph styles: {existing_styles}")
    
    # 查找所有可能的标题样式，包括自定义样式
    possible_h1_names = ['Heading 1', '标题 1', '自定义标题 1', '标题11', 'Heading 11', '标题1', 'Heading1']
    possible_h2_names = ['Heading 2', '标题 2', '自定义标题 2', '标题21', 'Heading 21', '标题2', 'Heading2']
    
    # 查找实际的标题样式
    actual_h1_name = None
    actual_h2_name = None
    
    for name in possible_h1_names:
        if name in existing_styles:
            actual_h1_name = name
            print(f"Found existing Heading 1 style: '{name}'")
            # 对现有样式进行字体设置
            try:
                heading1 = styles[name]
                heading1.font.bold = True
                heading1.font.name = '黑体'
                heading1.font.size = Pt(16)
                heading1.font.color.rgb = RGBColor(0, 0, 0)
                heading1.paragraph_format.outline_level = 0
                heading1.paragraph_format.keep_with_next = True
                
                # 强制设置XML级别的字体
                rPr = heading1.element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), '黑体')
                rFonts.set(qn('w:hAnsi'), '黑体')
                rFonts.set(qn('w:eastAsia'), '黑体')
                rFonts.set(qn('w:cs'), '黑体')
                rFonts.set(qn('w:hint'), 'eastAsia')
                print(f"Updated existing '{name}' style with black font")
            except Exception as e:
                print(f"Warning: Could not update existing style '{name}': {e}")
            break
    
    for name in possible_h2_names:
        if name in existing_styles:
            actual_h2_name = name
            print(f"Found existing Heading 2 style: '{name}'")
            # 对现有样式进行字体设置
            try:
                heading2 = styles[name]
                heading2.font.bold = True
                heading2.font.name = '宋体'
                heading2.font.size = Pt(10.5)
                heading2.font.color.rgb = RGBColor(0, 0, 0)
                heading2.paragraph_format.outline_level = 1
                
                # 强制设置XML级别的字体
                rPr = heading2.element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), '宋体')
                rFonts.set(qn('w:hAnsi'), '宋体')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:cs'), '宋体')
                rFonts.set(qn('w:hint'), 'eastAsia')
                print(f"Updated existing '{name}' style with SimSun font")
            except Exception as e:
                print(f"Warning: Could not update existing style '{name}': {e}")
            break
    
    # 如果没有找到，创建新的样式，但使用唯一名称
    if not actual_h1_name:
        try:
            print("Attempting to create Heading 1 style...")
            # 直接使用唯一名称避免冲突
            actual_h1_name = '自定义标题 1'
            heading1 = styles.add_style(actual_h1_name, 1)
            print(f"Created unique style: '{actual_h1_name}'")
            
            # 验证样式是否真的用期望的名称创建了
            created_styles = []
            for style in styles:
                if style.type.name == 'PARAGRAPH':
                    created_styles.append(style.name)
            print(f"Styles after creation: {created_styles}")
            
            # 查找基础样式
            base_styles = ['Normal', '正文', 'Body Text', 'Body']
            base_style = None
            for base_name in base_styles:
                if base_name in styles:
                    base_style = styles[base_name]
                    break
            
            if base_style:
                heading1.base_style = base_style
                print(f"Set '{actual_h1_name}' base style to '{base_style.name}'")
            
            heading1.font.bold = True
            heading1.font.name = '黑体'
            heading1.font.size = Pt(16)
            heading1.font.color.rgb = RGBColor(0, 0, 0)
            heading1.paragraph_format.outline_level = 0
            heading1.paragraph_format.keep_with_next = True
            print(f"Configured '{actual_h1_name}' style")
        except Exception as e:
            print(f"Warning: Could not create Heading 1 style: {e}")
    
    if not actual_h2_name:
        try:
            print("Attempting to create Heading 2 style...")
            # 直接使用唯一名称避免冲突
            actual_h2_name = '自定义标题 2'
            heading2 = styles.add_style(actual_h2_name, 1)
            print(f"Created unique style: '{actual_h2_name}'")
            
            # 验证样式是否真的用期望的名称创建了
            created_styles = []
            for style in styles:
                if style.type.name == 'PARAGRAPH':
                    created_styles.append(style.name)
            print(f"Styles after creation: {created_styles}")
            
            # 查找基础样式
            base_styles = ['Normal', '正文', 'Body Text', 'Body']
            base_style = None
            for base_name in base_styles:
                if base_name in styles:
                    base_style = styles[base_name]
                    break
            
            if base_style:
                heading2.base_style = base_style
                print(f"Set '{actual_h2_name}' base style to '{base_style.name}'")
            
            heading2.font.bold = True
            heading2.font.name = '宋体'  # 改为宋体
            heading2.font.size = Pt(10.5)
            heading2.font.color.rgb = RGBColor(0, 0, 0)
            heading2.paragraph_format.outline_level = 1
            
            # 强制设置所有字体类型为宋体
            rPr = heading2.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), '宋体')  # ASCII字符
            rFonts.set(qn('w:hAnsi'), '宋体')  # 高ANSI字符
            rFonts.set(qn('w:eastAsia'), '宋体')  # 东亚字符
            rFonts.set(qn('w:cs'), '宋体')  # 复杂脚本字符
            rFonts.set(qn('w:hint'), 'eastAsia')  # 字体提示
            print(f"Configured '{actual_h2_name}' style")
        except Exception as e:
            print(f"Warning: Could not create Heading 2 style: {e}")
    
    # 最终检查：重新获取实际的样式名称
    final_styles = []
    for style in styles:
        if style.type.name == 'PARAGRAPH':
            final_styles.append(style.name)
    
    print(f"Final paragraph styles: {final_styles}")
    
    # 重新查找实际的样式名称
    final_h1_name = None
    final_h2_name = None
    
    for name in possible_h1_names:
        if name in final_styles:
            final_h1_name = name
            break
    
    for name in possible_h2_names:
        if name in final_styles:
            final_h2_name = name
            break
    
    print(f"Final detected names: H1='{final_h1_name}', H2='{final_h2_name}'")
    
    # 如果仍然没有找到，说明可能是Word自动重命名了，我们需要通过样式属性来识别
    if not final_h1_name:
        print("No standard Heading 1 found, searching by outline level...")
        for style_name in final_styles:
            if '标题' in style_name and '1' in style_name:
                final_h1_name = style_name
                print(f"Found possible Heading 1 by pattern: '{style_name}'")
                break
    
    if not final_h2_name:
        print("No standard Heading 2 found, searching by outline level...")
        for style_name in final_styles:
            if '标题' in style_name and '2' in style_name:
                final_h2_name = style_name
                print(f"Found possible Heading 2 by pattern: '{style_name}'")
                break
    
    print(f"Final resolved names: H1='{final_h1_name}', H2='{final_h2_name}'")
    
    return final_h1_name, final_h2_name

def create_section_header_style(doc):
    """使用自定义标题2样式作为SectionHeader"""
    h1_name, h2_name = ensure_heading_styles_exist(doc)
    
    try:
        # 获取标题2样式
        if h2_name in doc.styles:
            heading2 = doc.styles[h2_name]
            print(f"Using '{h2_name}' style for SectionHeader")
        else:
            print("Error: No Heading 2 style available")
            return None
            
        # 自定义标题2的字体和格式
        try:
            heading2.font.bold = True
            heading2.font.name = '宋体'  # 改为宋体
            heading2.font.size = Pt(10.5)
            heading2.font.color.rgb = RGBColor(0, 0, 0)
            
            # 强制设置所有字体类型为宋体
            rPr = heading2.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), '宋体')  # ASCII字符
            rFonts.set(qn('w:hAnsi'), '宋体')  # 高ANSI字符
            rFonts.set(qn('w:eastAsia'), '宋体')  # 东亚字符
            rFonts.set(qn('w:cs'), '宋体')  # 复杂脚本字符
            rFonts.set(qn('w:hint'), 'eastAsia')  # 字体提示
            
            # 确保字体设置被应用
            if not hasattr(heading2.font, '_element'):
                heading2.font._element = rPr
            
            # 设置段落格式
            pPr = heading2.element.get_or_add_pPr()
            spacing = pPr.get_or_add_spacing()
            spacing.set(qn('w:after'), "0")
            spacing.set(qn('w:before'), "0")
            spacing.set(qn('w:line'), "240")
            spacing.set(qn('w:lineRule'), "auto")
            
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:left'), "0")
            ind.set(qn('w:right'), "0")
            ind.set(qn('w:firstLine'), "0")
            
            heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading2.paragraph_format.outline_level = 1
            
            print("Customized Heading 2 for SectionHeader use")
            return heading2
            
        except Exception as e:
            print(f"Warning: Error customizing Heading 2: {e}")
            return heading2
            
    except Exception as e:
        print(f"Error accessing Heading 2 style: {e}")
        return None

def create_route_header_style(doc):
    """使用自定义标题1样式作为RouteHeader"""
    h1_name, h2_name = ensure_heading_styles_exist(doc)
    
    try:
        # 获取标题1样式
        if h1_name in doc.styles:
            heading1 = doc.styles[h1_name]
            print(f"Using '{h1_name}' style for RouteHeader")
        else:
            print("Error: No Heading 1 style available")
            return None
            
        # 自定义标题1的字体和格式
        try:
            heading1.font.bold = True
            heading1.font.name = '黑体'
            heading1.font.size = Pt(16)
            heading1.font.color.rgb = RGBColor(0, 0, 0)
            
            # 强制设置所有字体类型为黑体
            rPr = heading1.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), '黑体')  # ASCII字符
            rFonts.set(qn('w:hAnsi'), '黑体')  # 高ANSI字符
            rFonts.set(qn('w:eastAsia'), '黑体')  # 东亚字符
            rFonts.set(qn('w:cs'), '黑体')  # 复杂脚本字符
            rFonts.set(qn('w:hint'), 'eastAsia')  # 字体提示
            
            # 确保字体设置被应用
            if not hasattr(heading1.font, '_element'):
                heading1.font._element = rPr
            
            # 设置段落格式
            pPr = heading1.element.get_or_add_pPr()
            spacing = pPr.get_or_add_spacing()
            spacing.set(qn('w:after'), "0")
            spacing.set(qn('w:before'), "0")
            spacing.set(qn('w:line'), "240")
            spacing.set(qn('w:lineRule'), "auto")
            
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:left'), "0")
            ind.set(qn('w:right'), "0")
            ind.set(qn('w:firstLine'), "0")
            
            heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading1.paragraph_format.outline_level = 0
            heading1.paragraph_format.keep_with_next = True
            
            print("Customized Heading 1 for RouteHeader use")
            return heading1
            
        except Exception as e:
            print(f"Warning: Error customizing Heading 1: {e}")
            return heading1
            
    except Exception as e:
        print(f"Error accessing Heading 1 style: {e}")
        return None

def create_footer(doc):
    """Add page number to footer (Center, SimSum, 9pt).
    
    Includes error handling for footer operations.
    """
    try:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        try:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception as e:
            print(f"Warning: Error setting footer font: {e}")
        
        try:
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
        except Exception as e:
            print(f"Warning: Error creating page number field: {e}")
    except Exception as e:
        print(f"Warning: Could not create footer: {e}")

def process_file(filename):
    print(f"Processing {filename}...")
    doc = docx.Document(filename)
    
    # Ensure heading styles exist and customize them (只调用一次)
    h1_name, h2_name = ensure_heading_styles_exist(doc)
    create_section_header_style(doc)
    create_route_header_style(doc)
    
    body = doc._body._element
    
    # Collect all children
    all_elements = [child for child in body.iterchildren()]
    
    # Identify blocks/chunks
    headers = ["点间路线描述", "分段路线上界线描述", "路线编号", "路线描述", "目标任务", "图幅编号", "地质点号", "路线小结", "路线自检"]
    
    # Keywords that need SectionHeader style (with Chinese colon)
    section_header_keywords = ["点上界线描述：", "点间路线描述：", "路线小结：", "路线自检："]
    
    chunks = []
    current_chunk = []
    current_header = "START"
    
    for elem in all_elements:
        text = get_text_from_element(elem).strip()
        
        # Skip empty lines (we will add them back where needed)
        # Check if it's truly empty (no table, no image, no sectPr)
        if not text and not elem.tag.endswith('tbl') and not elem.tag.endswith('sectPr') and not elem.findall('.//w:drawing', namespaces=elem.nsmap) and not elem.findall('.//w:pict', namespaces=elem.nsmap):
             continue
             

            
        is_header = False
        header_type = None
        
        for h in headers:
            if h in text:
                is_header = True
                header_type = h
                break
        
        if is_header:
            if current_chunk:
                chunks.append({"type": current_header, "elements": current_chunk})
            current_chunk = [elem]
            current_header = header_type
        else:
            current_chunk.append(elem)
            
    if current_chunk:
        chunks.append({"type": current_header, "elements": current_chunk})
        
    # Reorder Logic
    new_chunks = []
    
    # Extract Info for Header
    route_number = "Unknown"
    points = []
    
    for chunk in chunks:
        if chunk["type"] == "路线编号":
            text = get_text_from_element(chunk["elements"][0])
            match = re.search(r'[Ll]\d+', text)
            if match:
                route_number = match.group(0)
        if chunk["type"] == "地质点号":
            text = get_text_from_element(chunk["elements"][0])
            match = re.search(r'[Dd]\d+', text)
            if match:
                points.append(match.group(0))
                
    i = 0
    while i < len(chunks):
        chunk = chunks[i]
        
        if chunk["type"] == "点间路线描述":
            if i + 1 < len(chunks):
                if chunks[i+1]["type"] == "分段路线上界线描述":
                    # Swap
                    new_chunks.append(chunks[i+1])
                    new_chunks.append(chunk)
                    i += 2
                    continue
        
        new_chunks.append(chunk)
        i += 1
        
    # Clear body
    body.clear_content()
    
    # Rebuild body and apply formatting
    first_point_found = False
    
    for chunk in new_chunks:
        # Text Replacement
        if chunk["type"] == "分段路线上界线描述":
            # Insert empty line before Boundary
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), "240")
            spacing.set(qn('w:lineRule'), "auto")
            spacing.set(qn('w:after'), "0")
            pPr.append(spacing)
            p.append(pPr)
            body.append(p)

            if len(chunk["elements"]) > 0:
                # Replace with both possible colon formats (Chinese and English)
                replace_text_in_element(chunk["elements"][0], "分段路线上界线描述：", "点上界线描述：")
                replace_text_in_element(chunk["elements"][0], "分段路线上界线描述:", "点上界线描述：")
                replace_text_in_element(chunk["elements"][0], "分段路线上界线描述", "点上界线描述：")
            chunk["type"] = "界线描述"
            
        # Spacing Logic
        
        # 1. Before "地质点号" (Except first)
        if chunk["type"] == "地质点号":
            if first_point_found:
                # Insert empty line
                p = OxmlElement('w:p')
                body.append(p)
            first_point_found = True
            
        # 2. Before "点间路线描述" (Route)
        if chunk["type"] == "点间路线描述":
            p = OxmlElement('w:p')
            body.append(p)
            
        # Apply single spacing to all paragraphs in chunk
        format_chunk_spacing(chunk)
        
        for elem in chunk["elements"]:
            # Normalize all colons to Chinese colon
            normalize_colons_in_element(elem)
            
            # Check for "照片点坐标"
            text = get_text_from_element(elem)
            if "照片点坐标" in text:
                # Insert empty line before this element
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:line'), "240")
                spacing.set(qn('w:lineRule'), "auto")
                spacing.set(qn('w:after'), "0")
                pPr.append(spacing)
                p.append(pPr)
                body.append(p)
            
            # Ensure 0pt after for the element itself
            if elem.tag.endswith('p'):
                pPr = elem.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    elem.insert(0, pPr)
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                spacing.set(qn('w:after'), "0")
            
            body.append(elem)
    
    # Set Columns
    set_columns(doc, 2)
    
    # Format Fonts and apply SectionHeader style
    # 使用已获取的样式名称，不再重复调用
    for p in doc.paragraphs:
        # 先检查是否需要应用标题样式
        text = p.text.strip()
        style_applied = False
        
        for keyword in section_header_keywords:
            if keyword in text:
                if h2_name in doc.styles:
                    p.style = doc.styles[h2_name]
                    print(f"Applied '{h2_name}' style to section header: '{text}'")
                    
                    # 直接在每个run中设置字体，但保留大纲级别
                    for run in p.runs:
                        run.font.name = '宋体'
                        run.font.bold = True
                        run.font.size = Pt(10.5)
                        # 强制设置中文字体
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                    
                    # 如果段落没有runs，添加一个run并设置字体
                    if len(p.runs) == 0:
                        run = p.add_run(text)
                        run.font.name = '宋体'
                        run.font.bold = True
                        run.font.size = Pt(10.5)
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                    
                    # 确保大纲级别设置正确（用于导航功能）
                    pPr = p._element.get_or_add_pPr()
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.append(outlineLvl)
                    outlineLvl.set(qn('w:val'), "1")  # 二级标题
                    
                    style_applied = True
                break
        
        # 如果没有应用标题样式，则应用普通字体格式
        if not style_applied:
            format_paragraph_fonts(p)
        
    # Insert Header AFTER global formatting
    if points:
        start_point = points[0]
        end_point = points[-1]
        header_text = f"{route_number} ({start_point}-{end_point})"
        
        header_p = doc.paragraphs[0].insert_paragraph_before(header_text)
        
        # RouteHeader style已在前面创建，直接使用
        route_style = doc.styles[h1_name] if h1_name in doc.styles else None
        
        # Apply Heading 1 style to header paragraph (确保样式正确应用)
        if route_style:
            try:
                # 使用实际的样式名称
                if h1_name in doc.styles:
                    header_p.style = doc.styles[h1_name]
                    print(f"Applied '{h1_name}' style to header paragraph: '{header_text}'")
                    
                    # 直接在每个run中设置字体，但保留大纲级别
                    for run in header_p.runs:
                        run.font.name = '黑体'
                        run.font.bold = True
                        run.font.size = Pt(16)
                        # 强制设置中文字体
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
                    
                    # 如果段落没有runs，添加一个run并设置字体
                    if len(header_p.runs) == 0:
                        run = header_p.add_run(header_text)
                        run.font.name = '黑体'
                        run.font.bold = True
                        run.font.size = Pt(16)
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
                    
                    # 确保大纲级别设置正确（用于导航功能）
                    pPr = header_p._element.get_or_add_pPr()
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.append(outlineLvl)
                    outlineLvl.set(qn('w:val'), "0")  # 一级标题
                    
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:before'), "0")
                    spacing.set(qn('w:after'), "0")
                    spacing.set(qn('w:line'), "240")
                    spacing.set(qn('w:lineRule'), "auto")
                    pPr.append(spacing)
                    
                    # 设置大纲级别确保导航功能
                    outlineLvl = OxmlElement('w:outlineLvl')
                    outlineLvl.set(qn('w:val'), "0")  # 一级标题
                    pPr.append(outlineLvl)
                    
            except Exception as e:
                print(f"Warning: Could not apply Heading 1 style to paragraph: {e}")
        else:
            print(f"Warning: RouteHeader style creation failed, using default formatting")
            
    # Resize Images
    resize_images(doc)

    # Create Footer with Page Number
    create_footer(doc)
    
    new_filename = filename.replace(".docx", "_formatted.docx")
    doc.save(new_filename)
    print(f"Saved to {new_filename}")

def run_batch():
    """Run the formatting process for all matching files from DGSS导出报告 folder."""
    import glob
    import sys
    
    # Determine if application is a script file or frozen exe
    if getattr(sys, 'frozen', False):
        script_dir = os.path.dirname(sys.executable)
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 输入文件夹：DGSS导出报告
    input_dir = os.path.join(script_dir, "DGSS导出报告")
    # 输出文件夹：重新排版的报告
    output_dir = os.path.join(script_dir, "重新排版的报告")
    
    print("=" * 60)
    print("Word文档格式化工具")
    print("=" * 60)
    print(f"输入文件夹: {input_dir}")
    print(f"输出文件夹: {output_dir}")
    
    # 检查输入文件夹是否存在
    if not os.path.exists(input_dir):
        print(f"\n错误：输入文件夹不存在: {input_dir}")
        print("请在word文件夹下创建'DGSS导出报告'文件夹，并将原始文档放入其中。")
        return
    
    # 创建输出文件夹（如果不存在）
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出文件夹: {output_dir}")
    
    # 获取输入文件夹中的所有docx文件
    all_docx_files = glob.glob(os.path.join(input_dir, "*.docx"))
    
    # 筛选条件：
    # 1. 不是临时文件（以~$开头）
    # 2. 必须符合路线号格式（L[数字].docx）
    docx_files = []
    for f in all_docx_files:
        basename = os.path.basename(f)
        if basename.startswith("~$"):
            continue
        # 检查是否符合路线号格式：L + 数字 + .docx
        if re.match(r'^L\d+\.docx$', basename):
            docx_files.append(f)
    
    # 按路线号排序
    docx_files.sort()
    
    print(f"\n找到 {len(docx_files)} 个路线文件待处理\n")
    
    processed_count = 0
    error_count = 0
    
    for input_path in docx_files:
        basename = os.path.basename(input_path)
        output_filename = basename.replace(".docx", "_formatted.docx")
        output_path = os.path.join(output_dir, output_filename)
        
        try:
            print(f"处理: {basename}")
            
            # 处理文件（需要修改process_file函数接受输入和输出路径）
            doc = docx.Document(input_path)
            
            # 执行原有的格式化逻辑（从process_file中提取）
            # Ensure heading styles exist and customize them
            h1_name, h2_name = ensure_heading_styles_exist(doc)
            create_section_header_style(doc)
            create_route_header_style(doc)
            
            body = doc._body._element
            
            # Collect all children
            all_elements = [child for child in body.iterchildren()]
            
            # Identify blocks/chunks
            headers = ["点间路线描述", "分段路线上界线描述", "路线编号", "路线描述", "目标任务", "图幅编号", "地质点号", "路线小结", "路线自检"]
            
            # Keywords that need SectionHeader style (with Chinese colon)
            section_header_keywords = ["点上界线描述：", "点间路线描述：", "路线小结：", "路线自检："]
            
            chunks = []
            current_chunk = []
            current_header = "START"
            
            for elem in all_elements:
                text = get_text_from_element(elem).strip()
                
                # Skip empty lines
                if not text and not elem.tag.endswith('tbl') and not elem.tag.endswith('sectPr') and not elem.findall('.//w:drawing', namespaces=elem.nsmap) and not elem.findall('.//w:pict', namespaces=elem.nsmap):
                     continue
                
                is_header = False
                header_type = None
                
                for h in headers:
                    if h in text:
                        is_header = True
                        header_type = h
                        break
                
                if is_header:
                    if current_chunk:
                        chunks.append({"type": current_header, "elements": current_chunk})
                    current_chunk = [elem]
                    current_header = header_type
                else:
                    current_chunk.append(elem)
                    
            if current_chunk:
                chunks.append({"type": current_header, "elements": current_chunk})
                
            # Reorder Logic
            new_chunks = []
            
            # Extract Info for Header
            route_number = "Unknown"
            points = []
            
            for chunk in chunks:
                if chunk["type"] == "路线编号":
                    text = get_text_from_element(chunk["elements"][0])
                    match = re.search(r'[Ll]\d+', text)
                    if match:
                        route_number = match.group(0)
                if chunk["type"] == "地质点号":
                    text = get_text_from_element(chunk["elements"][0])
                    match = re.search(r'[Dd]\d+', text)
                    if match:
                        points.append(match.group(0))
                        
            i = 0
            while i < len(chunks):
                chunk = chunks[i]
                
                if chunk["type"] == "点间路线描述":
                    if i + 1 < len(chunks):
                        if chunks[i+1]["type"] == "分段路线上界线描述":
                            # Swap
                            new_chunks.append(chunks[i+1])
                            new_chunks.append(chunk)
                            i += 2
                            continue
                
                new_chunks.append(chunk)
                i += 1
                
            # Clear body
            body.clear_content()
            
            # Rebuild body and apply formatting
            first_point_found = False
            
            for chunk in new_chunks:
                # Text Replacement
                if chunk["type"] == "分段路线上界线描述":
                    # Insert empty line before Boundary
                    p = OxmlElement('w:p')
                    pPr = OxmlElement('w:pPr')
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:line'), "240")
                    spacing.set(qn('w:lineRule'), "auto")
                    spacing.set(qn('w:after'), "0")
                    pPr.append(spacing)
                    p.append(pPr)
                    body.append(p)

                    if len(chunk["elements"]) > 0:
                        # Replace with both possible colon formats
                        replace_text_in_element(chunk["elements"][0], "分段路线上界线描述：", "点上界线描述：")
                        replace_text_in_element(chunk["elements"][0], "分段路线上界线描述:", "点上界线描述：")
                        replace_text_in_element(chunk["elements"][0], "分段路线上界线描述", "点上界线描述：")
                    chunk["type"] = "界线描述"
                    
                # Spacing Logic
                
                # 1. Before "地质点号" (Except first)
                if chunk["type"] == "地质点号":
                    if first_point_found:
                        # Insert empty line
                        p = OxmlElement('w:p')
                        body.append(p)
                    first_point_found = True
                    
                # 2. Before "点间路线描述" (Route)
                if chunk["type"] == "点间路线描述":
                    p = OxmlElement('w:p')
                    body.append(p)
                    
                # Apply single spacing to all paragraphs in chunk
                format_chunk_spacing(chunk)
                
                for elem in chunk["elements"]:
                    # Normalize all colons to Chinese colon
                    normalize_colons_in_element(elem)
                    
                    # Check for "照片点坐标"
                    text = get_text_from_element(elem)
                    if "照片点坐标" in text:
                        # Insert empty line before this element
                        p = OxmlElement('w:p')
                        pPr = OxmlElement('w:pPr')
                        spacing = OxmlElement('w:spacing')
                        spacing.set(qn('w:line'), "240")
                        spacing.set(qn('w:lineRule'), "auto")
                        spacing.set(qn('w:after'), "0")
                        pPr.append(spacing)
                        p.append(pPr)
                        body.append(p)
                    
                    # Ensure 0pt after for the element itself
                    if elem.tag.endswith('p'):
                        pPr = elem.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            elem.insert(0, pPr)
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:after'), "0")
                    
                    body.append(elem)
            
            # Set Columns
            set_columns(doc, 2)
            
            # Format Fonts and apply SectionHeader style
            for p in doc.paragraphs:
                # 先检查是否需要应用标题样式
                text = p.text.strip()
                style_applied = False
                
                for keyword in section_header_keywords:
                    if keyword in text:
                        if h2_name in doc.styles:
                            p.style = doc.styles[h2_name]
                            
                            # 直接在每个run中设置字体
                            for run in p.runs:
                                run.font.name = '宋体'
                                run.font.bold = True
                                run.font.size = Pt(10.5)
                                if run._element.rPr is None:
                                    run._element._add_rPr()
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                            
                            if len(p.runs) == 0:
                                run = p.add_run(text)
                                run.font.name = '宋体'
                                run.font.bold = True
                                run.font.size = Pt(10.5)
                                if run._element.rPr is None:
                                    run._element._add_rPr()
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                            
                            # 确保大纲级别设置正确
                            pPr = p._element.get_or_add_pPr()
                            outlineLvl = pPr.find(qn('w:outlineLvl'))
                            if outlineLvl is None:
                                outlineLvl = OxmlElement('w:outlineLvl')
                                pPr.append(outlineLvl)
                            outlineLvl.set(qn('w:val'), "1")
                            
                            style_applied = True
                        break
                
                if not style_applied:
                    format_paragraph_fonts(p)
                
            # Insert Header
            if points:
                start_point = points[0]
                end_point = points[-1]
                header_text = f"{route_number} ({start_point}-{end_point})"
                
                header_p = doc.paragraphs[0].insert_paragraph_before(header_text)
                
                if h1_name in doc.styles:
                    header_p.style = doc.styles[h1_name]
                    
                    for run in header_p.runs:
                        run.font.name = '黑体'
                        run.font.bold = True
                        run.font.size = Pt(16)
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
                    
                    if len(header_p.runs) == 0:
                        run = header_p.add_run(header_text)
                        run.font.name = '黑体'
                        run.font.bold = True
                        run.font.size = Pt(16)
                        if run._element.rPr is None:
                            run._element._add_rPr()
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
                    
                    pPr = header_p._element.get_or_add_pPr()
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.append(outlineLvl)
                    outlineLvl.set(qn('w:val'), "0")
                    
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:before'), "0")
                    spacing.set(qn('w:after'), "0")
                    spacing.set(qn('w:line'), "240")
                    spacing.set(qn('w:lineRule'), "auto")
                    pPr.append(spacing)
                    
            # Resize Images
            resize_images(doc)

            # Create Footer with Page Number
            create_footer(doc)
            
            # 保存到输出文件夹
            doc.save(output_path)
            print(f"  ✓ 已保存: {output_filename}\n")
            processed_count += 1
            
        except Exception as e:
            print(f"  ✗ 错误: {e}\n")
            error_count += 1
            continue
    
    print("=" * 60)
    print("处理完成!")
    print(f"成功处理: {processed_count} 个文件")
    print(f"失败: {error_count} 个文件")
    print(f"输出位置: {output_dir}")
    print("=" * 60)

if __name__ == "__main__":
    run_batch()
